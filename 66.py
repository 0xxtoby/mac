import random
import re

import docx
import zipfile
import os
import shutil

'''读取word中的文本'''

def teat():
    import re
    a = """
第四十八讲
如何正确地使用法宝




    """
    print(re.search("([0-9]*)([a-z]*)([0-9]*)", a).group())  # 123abc456,返回整体
    print(re.search("([0-9]*)([a-z]*)([0-9]*)", a).group(1))  # 123
    print(re.search("([0-9]*)([a-z]*)([0-9]*)", a).group(2))  # abc
    print(re.search("([0-9]*)([a-z]*)([0-9]*)", a).group(3))  # 456
    print(re.search("第.*讲", a))  # 456

def gettxt():
    file = docx.Document("data/4.docx")
    print("段落数:" + str(len(file.paragraphs)))  # 段落数为13，每个回车隔离一段

    # 输出每一段的内容
    # for para in file.paragraphs:
    #     print(para.text)

    # 输出段落编号及段落内容
    ii=0
    sum=20

    for i in range(len(file.paragraphs)):
        if len(file.paragraphs[i].text.replace(' ', '')) > 4:
            a=""
            a=str(file.paragraphs[i].text)
            # print(a)
            # print(re.search("第.*讲",a))
            if(re.search("第.讲",a)==None):
                ram=random.randint(0,9)
                if (ram>5 and len(file.paragraphs[i].text.replace(' ', '')) > 30):
                    ii = ii + 1
                    if(ram>3):
                     sum=sum+1
                    print(" "+ str(ii) + "：" + file.paragraphs[i].text+"("+str(sum)+")")
                    print(" ")

            else:
                # print("++++++++++++++"+"第" + str(i) + "段的内容是：" + file.paragraphs[i].text)
                print(re.search("第.*讲",a).group(0))
                print(" ")






if __name__ == '__main__':
    # teat()
    gettxt()