# -*- coding:utf-8 -*-

import random
import re

import docx
import zipfile
import os
import shutil

'''读取word中的文本'''

def re_tat(sss):
    s = '''欣赏婉约词的门径 北宋名家词浅释 无名氏（一首） 菩萨蛮（平林漠漠烟如织） 范仲淹（一首） 渔家傲（塞下秋来风景异） 张先（三首） 一丛花令（伤高怀远几时穷） 天仙子（水调数声持酒听） 醉垂鞭（双蝶绣罗裙） 晏殊（二首） 蝶恋花（槛菊愁烟兰泣露） 破阵子（燕子来时新社） 欧阳修（一首） 踏莎行（候馆梅残） 柳永（七首） 雨霖铃（寒蝉凄切） 曲玉管（陇首云飞） 夜半乐（冻云黯淡天气） 卜算子慢（江枫渐老） 安公子（远岸收残雨） 八声甘州（对潇潇暮雨洒江天） 望海潮（东南形胜） 晏几道（六首） 蝶恋花（醉别西楼醒不记） 阮郎归（天边金掌露成霜） 鹧鸪天（小令尊前见玉箫） 临江仙（梦后楼台高锁） 鹧鸪天（彩袖殷勤捧玉钟） 浣溪沙（日日双眉斗画长） 苏轼（二首） 水调歌头（明月几时有） 念奴娇（大江东去） 秦观（六首） 八六子（倚危亭） 满庭芳（山抹微云） 浣溪沙（漠漠轻寒上小楼） 望海潮（梅英疏淡） 满庭芳（晓色云开） 鹊桥仙（纤云弄巧） 贺铸（四首） 芳心苦（杨柳回塘） 横塘路（凌波不过横塘路） 薄幸（淡妆多态） 将进酒（城下路） 周邦彦（七首） 瑞龙吟（章台路） 兰陵王（柳阴直） 夜飞鹊（河桥送人处） 玉楼春（桃溪不作从容住） 解连环（怨怀无托） 拜星月慢（夜色催更） 过秦楼（水浴清蟾） 李清照（五首） 凤凰台上忆吹箫（香冷金猊） 念奴娇（萧条庭院） 声声慢（寻寻觅觅） 武陵春（风住尘香花已尽） 永遇乐（落日熔金） 姜夔词小札 小重山令（人绕湘皋月坠时） 江梅引（人间离别易多时） 点绛唇（燕雁无心） 鹧鸪天（京洛风流绝代人） 鹧鸪天（巷陌风光纵赏时） 鹧鸪天（肥水东流无尽期） 踏莎行（燕燕轻盈） 浣溪沙（著酒行行满袂风） 浣溪沙（雁怯重云不肯啼） 霓裳中序第一（亭皋正望极） 齐天乐（庾郎先自吟愁赋） 一萼红（古城阴） 念奴娇（闹红一舸） 月下笛（与客携壶） 琵琶仙（双桨来时） 玲珑四犯（叠鼓夜寒） 扬州慢（淮左名都） 长亭怨慢（渐吹尽、枝头香絮） 淡黄柳（空城晓角） 暗香（旧时月色） 疏影（苔枝缀玉） 张炎词小札 南浦（波暖绿粼粼） 解连环（楚江空晚） 高阳台（接叶巢莺） 高阳台（古木迷鸦） 扫花游（嫩寒禁暖） 渡江云（山空天入海） 渡江云（锦香缭绕地） 声声慢（寒花清事） 声声慢（平沙催晓） 声声慢（烟堤小舫） 声声慢（山风古道） 绮罗香（万里飞霜） 壶中天（扬舲万里） 八声甘州（记玉关踏雪事清游） 八声甘州（望涓涓一水隐芙蓉） 台城路（朗吟未了西湖酒） 台城路（十年前事翻疑梦） 台城路（春风不暖垂杨树） 忆旧游（叹江潭树老） 满庭芳（晴皎霜花） 凄凉犯（萧疏野柳鸣寒雨） 后记 出版说明 无名氏（一首） 

      '''
    s_list = []
    for ss in re.findall("\S*", s):
        if ss != "":
            s_list.append(ss)

        if  ss==sss:
            return True
    # print(s_list)
    return False


def teat():
    import re





    a = """
第四十八讲
如何正确地使用法宝




    """
    print(re.search("([0-9]*)([a-z]*)([0-9]*)", a).group())  # 123abc456,返回整体
    print(re.search("([0-9]*)([a-z]*)([0-9]*)", a).group(1))  # 123
    print(re.search("([0-9]*)([a-z]*)([0-9]*)", a).group(2))  # abc
    print(re.search("([0-9]*)([a-z]*)([0-9]*)", a).group(3))  # 456
    print(re.search("第.*讲", a))  # 456







def gettxt():
    file = docx.Document("data/8.docx")
    print("段落数:" + str(len(file.paragraphs)))  # 段落数为13，每个回车隔离一段

    # 输出每一段的内容
    # for para in file.paragraphs:
    #     print(para.text)

    # 输出段落编号及段落内容
    ii=0
    sum=20

    for i in range(len(file.paragraphs)):
        if len(file.paragraphs[i].text.replace(' ', '')) >4:
            a=""
            a=str(file.paragraphs[i].text)
            # print(a)
            # print(" ++++++++++++++++")
            # print(re.search("第.*讲",a))
            if(re.search("第...讲",a)==None):
                ram=random.randint(0,9)
                if (ram>3 and len(file.paragraphs[i].text.replace(' ', '')) > 30):
                    ii = ii + 1
                    if(ram>3):
                     sum=sum+1
                    print(" "+ str(ii) + "：" + file.paragraphs[i].text+"("+str(sum)+")")
                    print(" ")
                else:
                    # print(5)
                    pass

            else:
                # print("++++++++++++++"+"第" + str(i) + "段的内容是：" + file.paragraphs[i].text)
                print(a)
                print(" ")






if __name__ == '__main__':
    # teat()
    gettxt()
    # re_tat("无名氏（一首）")