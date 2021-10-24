import random
import re

import docx
import zipfile
import os
import shutil

'''读取word中的文本'''

def teat():
    import re
    a = """
第四十八讲
如何正确地使用法宝




    """
    print(re.search("([0-9]*)([a-z]*)([0-9]*)", a).group())  # 123abc456,返回整体
    print(re.search("([0-9]*)([a-z]*)([0-9]*)", a).group(1))  # 123
    print(re.search("([0-9]*)([a-z]*)([0-9]*)", a).group(2))  # abc
    print(re.search("([0-9]*)([a-z]*)([0-9]*)", a).group(3))  # 456
    print(re.search("第.*讲", a))  # 456

def gettxt():
    file = docx.Document("6666.docx")
    print("段落数:" + str(len(file.paragraphs)))  # 段落数为13，每个回车隔离一段

    # 输出每一段的内容
    # for para in file.paragraphs:
    #     print(para.text)

    # 输出段落编号及段落内容
    strrr="""编前语
三国史话
楔子
宦官
外戚
黄巾
历史和文学
后汉的地理
董卓的扰乱
曹操是怎样强起来的
曹孟德移驾幸许都
袁绍和曹操的战争
赤壁之战的真相
刘备取益州和孙权取荆州
替魏武帝辨诬
从曹操到司马懿
替魏延辨诬
姜维和钟会
史话之余
孙吴为什么要建都南京
司马懿如何人
司马氏之兴亡
晋代豪门斗富
其他
诸葛亮南征考
诸葛亮随身衣食悉仰于官不别治生
诸葛亮治戎
如其不才君可自取
奖率三军臣职是当
袁曹成败
论魏武帝
曹嵩之死
魏时将帅之骄
魏太祖征乌丸
关羽欲杀曹公
孙策欲袭许
孙氏父子轻佻
姜维不速救成都
用人以抚绥新附
君与王之别
兵无铠甲
文臣轻视军人
张纯之叛
边章、韩遂
李邈
马钧
罢社
吞泥
三国之校事
    """



    ii=0
    sum=0

    for i in range(len(file.paragraphs)):
        if len(file.paragraphs[i].text.replace(' ', '')) > 1:
            a=""
            a=str(file.paragraphs[i].text)
            # print(a)
            # print(re.search("第.节",a))
            if(re.search(a,strrr)!=None):

                print(a)
                # print(re.search("第.节", a).group(0))
                print(" ")



            else:
                ram=random.randint(0,9)
                if (ram>5 and len(file.paragraphs[i].text.replace(' ', '')) > 20):
                    ii = ii + 1
                    if(ram>3):
                     sum=sum+1
                    print(" "+ str(ii) + "：" + file.paragraphs[i].text+"("+str(sum)+")")
                    print(" ")

            # print("++++++++++++++"+"第" + str(i) + "段的内容是：" + file.paragraphs[i].text)


if __name__ == '__main__':
    teat()
    gettxt()